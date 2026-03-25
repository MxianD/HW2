#!/usr/bin/env python3
"""Convert docs/HW2_Writeup.md to docs/HW2_Writeup.docx (Word). Run from repo root: python docs/md_to_docx.py"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
DEFAULT_MD = ROOT / "HW2_Writeup.md"
DEFAULT_DOCX = ROOT / "HW2_Writeup.docx"


def add_formatted_paragraph(doc: Document, text: str) -> None:
    if not text.strip():
        return
    p = doc.add_paragraph()
    pos = 0
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`\n]+`)")
    for m in pattern.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos : m.start()])
        chunk = m.group(1)
        if chunk.startswith("**"):
            r = p.add_run(chunk[2:-2])
            r.bold = True
        else:
            r = p.add_run(chunk[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    i = 0
    if lines and lines[0].strip() == "---":
        i = 1
        while i < len(lines) and lines[i].strip() != "---":
            i += 1
        i += 1

    para_buf: list[str] = []

    def flush_para() -> None:
        nonlocal para_buf
        if not para_buf:
            return
        content = " ".join(para_buf).strip()
        if content:
            add_formatted_paragraph(doc, content)
        para_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_para()
            fence_rest = stripped[3:].strip()
            if fence_rest.startswith("{=latex}"):
                block: list[str] = []
                i += 1
                while i < len(lines) and lines[i].strip() != "```":
                    block.append(lines[i])
                    i += 1
                i += 1
                if any("\\newpage" in ln for ln in block):
                    doc.add_page_break()
                continue

            block = []
            i += 1
            while i < len(lines) and lines[i].strip() != "```":
                block.append(lines[i])
                i += 1
            i += 1
            code_text = "\n".join(block)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        if stripped.startswith("### "):
            flush_para()
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("## "):
            flush_para()
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue

        if stripped == "---":
            flush_para()
            i += 1
            continue

        if stripped == "":
            flush_para()
            i += 1
            continue

        para_buf.append(line.rstrip())
        i += 1

    flush_para()
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    md = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_MD
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_DOCX
    if not md.is_file():
        print(f"Missing {md}", file=sys.stderr)
        sys.exit(1)
    convert(md, out)
